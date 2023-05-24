import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
from PyQt6.QtGui import QIcon
from PyQt6.QtCore import QTimer, QThread, pyqtSignal
import jieba
from wordcloud import WordCloud, ImageColorGenerator
import numpy as np
from PIL import Image
import time
import os
import chardet
from docx import Document
from new_cloud import Ui_wordcloud


class WorkerThread(QThread):
    my_str = pyqtSignal(str)

    def __init__(self, txt_path, wallpaper_path, out_path, stopword, font, backcolor, switch_stopword_type):
        QThread.__init__(self)
        self.switch_stopword_isChecked = switch_stopword_type
        self.txt_path = txt_path
        self.wallpaper_path = wallpaper_path
        self.out_path = out_path
        self.stopword = stopword
        self.font = font
        self.backcolor = backcolor

    def is_docx_file(self, path):
        return path.lower().endswith(".docx")

    def read_docx(self):
        try:
            doc = Document(self.txt_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except:
            self.my_str.emit("error_docx")

    def read_stopword_docx(self):
        doc = Document(self.stopword)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content.splitlines()

    def is_file_txt(self):
        with open(self.txt_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding_txt = result['encoding']
        with open(self.txt_path, 'r', encoding=encoding_txt) as f:
            text = f.read()
        return text

    def is_stopword_txt(self):
        with open(self.stopword, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding_stopword = result['encoding']
        with open(self.stopword, 'r', encoding=encoding_stopword) as f:
            stop_words = f.read().splitlines()
        return stop_words

    def word(self):
        try:
            if self.txt_path != "" and self.wallpaper_path != "" and self.out_path != "":
                file_type = self.is_docx_file(self.txt_path)
                stopword_type = self.is_docx_file(self.stopword)
                if file_type:
                    text = self.read_docx()
                    if stopword_type:
                        stopword = self.read_stopword_docx()
                        self.product_cloud(text, stopword)
                    else:
                        stopword = self.is_stopword_txt()
                        self.product_cloud(text, stopword)
                else:
                    text = self.is_file_txt()
                    if stopword_type:
                        stopword = self.read_stopword_docx()
                        self.product_cloud(text, stopword)
                    else:
                        stopword = self.is_stopword_txt()
                        self.product_cloud(text, stopword)
            else:
                self.my_str.emit("error_path")
        except:
            self.my_str.emit("error_font")

    def definl_nostopword(self):
        if self.is_docx_file(self.txt_path):
            text = self.read_docx()
            stopword = self.is_stopword_txt()
            self.product_cloud(text, stopword)
        else:
            text = self.is_file_txt()
            stopword = self.is_stopword_txt()
            self.product_cloud(text, stopword)

    def product_cloud(self, text, stop_words):
        print(text)
        words = jieba.cut(text)
        words = [word for word in words if word not in stop_words]
        mask = np.array(Image.open(self.wallpaper_path))
        color_func = ImageColorGenerator(mask)
        wc = WordCloud(font_path=self.font, background_color=self.backcolor, max_words=2000, mask=mask,
                       color_func=color_func)
        wc.generate(' '.join(words))
        timestamp = time.time()
        formatted_time = time.strftime("%m-%d-%H-%M-%S", time.localtime(timestamp))
        wc.to_file(self.out_path + "/" + formatted_time + ".png")
        self.my_str.emit("success")
    def run(self):
        try:
            if self.stopword != "./default_stopword.txt":
                if self.switch_stopword_isChecked:
                    self.word()
                else:
                    self.stopword = "./null.txt"
                    self.word()
            else:
                if self.switch_stopword_isChecked:
                    self.definl_nostopword()
                else:
                    self.stopword = "./null.txt"
                    self.word()
        except:
            self.my_str.emit("error")


class MyApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.start_time = None
        self.worker_thread = None
        self.ui = Ui_wordcloud()
        self.ui.setupUi(self)
        self.setFixedSize(400, 300)
        self.ui.running_time_label.setText("Running Time: 00:00:00")
        self.ui.pushButton.clicked.disconnect()
        self.ui.pushButton.clicked.connect(self.run)

        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_running_time)

        self.ui.inputpath_file.clicked.disconnect()
        self.ui.inputpath_file.clicked.connect(self.open)

        self.ui.inputpath.editingFinished.disconnect()
        self.ui.inputpath.editingFinished.connect(self.inputpath_txt)

        self.ui.input_wallpaper.clicked.disconnect()
        self.ui.input_wallpaper.clicked.connect(self.select_wallpaper_file)

        self.ui.pactuer_path.editingFinished.disconnect()
        self.ui.pactuer_path.editingFinished.connect(self.inputpath_wallpaper)

        self.ui.outputpath_file.clicked.disconnect()
        self.ui.outputpath_file.clicked.connect(self.openpath)

        self.ui.output_path.editingFinished.disconnect()
        self.ui.output_path.editingFinished.connect(self.inputpath_out)

        self.ui.stopword_file.clicked.disconnect()
        self.ui.stopword_file.clicked.connect(self.open_stopword)

        self.ui.stopword_path.editingFinished.disconnect()
        self.ui.stopword_path.editingFinished.connect(self.input_stopword_path)

        self.font_name = {"simsun.ttc": "宋体", "msyh.ttc": "微软雅黑",
                          "W02.ttf": "站酷仓耳渔阳", "LXGWWenKaiMono-Light.ttf": "霞鹜文楷(手写)"}
        self.font_path = "./font"
        self.list_font()
        self.backcolor_list()
        self.txt_path = ""
        self.wallpaper_path = ""
        self.out_path = ""
        self.stopword = "./default_stopword.txt"
        self.font = ''
        self.backcolor = ''
        self.switch_stopword_type = True

    def list_font(self):
        font_list = os.listdir(self.font_path)
        for font in font_list:
            try:
                font_name = self.font_name[font]
                self.ui.select_font.addItem(font_name)
            except KeyError:
                self.ui.select_font.addItem(font.split(".")[0])
                self.font_name[font] = font
        self.ui.select_font.setCurrentText("宋体")

    def backcolor_list(self):
        color_list = ["white", "black", "pink", "yellow"]
        for backcolor in color_list:
            self.ui.select_backbox.addItem(backcolor)

    def select_backcolor(self):
        self.backcolor = self.ui.select_backbox.currentText()

    def selected_font(self):
        self.font = self.font_path + "/" + self.get_keys_from_value(self.font_name, self.ui.select_font.currentText())

    def get_keys_from_value(self, dictionary, value):
        keys = []
        for key, val in dictionary.items():
            if val == value:
                keys.append(key)
        return keys[0]

    def input_stopword_path(self):
        self.stopword = self.ui.stopword_path.text()

    def inputpath_txt(self):
        self.txt_path = self.ui.inputpath.text()

    def inputpath_wallpaper(self):
        self.wallpaper_path = self.ui.pactuer_path.text()

    def inputpath_out(self):
        self.out_path = self.ui.output_path.text()

    def openpath(self):
        file_dialog = QFileDialog()
        file_path = file_dialog.getExistingDirectoryUrl(self, "选择文件路径")
        self.ui.output_path.setText(file_path.toLocalFile())
        self.inputpath_out()

    def open(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "选择输入文件")
        self.ui.inputpath.setText(file_path)
        self.inputpath_txt()

    def select_wallpaper_file(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "选择输入文件")
        self.ui.pactuer_path.setText(file_path)
        self.inputpath_wallpaper()

    def open_stopword(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "选择输入文件")
        self.ui.stopword_path.setText(file_path)
        self.input_stopword_path()

    def update_running_time(self):
        current_time = time.time()
        running_time = current_time - self.start_time
        formatted_time = time.strftime("%H:%M:%S", time.gmtime(running_time))
        self.ui.running_time_label.setText("Running Time: " + formatted_time)

    def worker_finished(self, out_str):
        self.timer.stop()
        if out_str == "success":
            self.message("提示", "成功！")
        elif out_str == "error":
            self.message("提示", "请检查输入文件!!")
        elif out_str == "error_path":
            self.message("提示", "请输入相应路径!")
        elif out_str == "error_font":
            self.message("提示", "字体不兼容!")
        elif out_str == "error_docx":
            self.message("提示！", "请检查输入文件docx")

    def message_box_clicked(self):
        self.ui.pushButton.setEnabled(True)
        self.ui.switch_stopword.setEnabled(True)

    def message(self, title, content):
        message_box = QMessageBox()
        message_box.setWindowTitle(title)  # Set title
        message_box.setText(content)  # Set text content
        message_box.setIcon(QMessageBox.Icon.Information)  # Set icon type
        message_box.finished.connect(self.message_box_clicked)
        message_box.exec()

    def run(self):
        self.ui.running_time_label.setText("Running Time: 00:00:00")
        self.selected_font()
        self.select_backcolor()

        self.start_time = time.time()
        self.timer.start(1000)
        self.ui.pushButton.setEnabled(False)
        self.ui.switch_stopword.setEnabled(False)
        print(self.stopword)
        print(self.font)
        print(self.backcolor)
        print(self.switch_stopword_type)
        print(self.txt_path)
        print(self.wallpaper_path)
        self.worker_thread = WorkerThread(txt_path=self.txt_path, wallpaper_path=self.wallpaper_path,
                                          out_path=self.out_path, stopword=self.stopword, font=self.font,

                                          backcolor=self.backcolor, switch_stopword_type=self.switch_stopword_type)
        self.worker_thread.my_str.connect(self.worker_finished)
        self.worker_thread.start()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MyApp()
    app_icon = QIcon("./wordcloud.ico")
    app.setWindowIcon(app_icon)
    window.show()
    sys.exit(app.exec())
